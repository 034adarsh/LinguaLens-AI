import streamlit as st
from io import StringIO
from reportlab.pdfgen import canvas
from transformers import MarianMTModel, MarianTokenizer
import os
import fitz  # PyMuPDF
import mimetypes
import openpyxl
import docx
import csv
import logging
import sys
import traceback
import tempfile
import sentencepiece  # Explicitly import to ensure availability

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    stream=sys.stdout,
    force=True
)
logger = logging.getLogger(__name__)

# Log SentencePiece version for debugging
logger.info(f"SentencePiece version: {sentencepiece.__version__}")

# Create output directory
OUTPUT_DIR = "translated_files"
os.makedirs(OUTPUT_DIR, exist_ok=True)
logger.info(f"Created output directory: {OUTPUT_DIR}")

# Global model cache
model_cache = {}

# Supported languages
LANGUAGES = {
    "English": "en",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Italian": "it",
    "Hindi": "hi",
    "Korean": "ko",
    "Japanese": "ja",
    "Chinese": "zh",
    "Arabic": "ar",
    "Russian": "ru",
    "Portuguese": "pt",
}

# MIME types for download
MIME_TYPES = {
    "txt": "text/plain",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "csv": "text/csv",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
}

# Utility Functions

def detect_file_type(file_path):
    """Detect file type based on MIME or extension."""
    mime_type, _ = mimetypes.guess_type(file_path)
    if mime_type and "pdf" in mime_type:
        return ".pdf"
    for ext in [".docx", ".xlsx", ".csv", ".txt"]:
        if file_path.endswith(ext):
            return ext
    logger.error(f"Unsupported file type: {file_path}")
    raise ValueError("Unsupported file type")

def chunk_text(text, max_length=512):
    """Split text into chunks for translation."""
    sentences = text.split(". ")
    chunks, chunk = [], ""
    for sentence in sentences:
        if len(chunk) + len(sentence) + 2 < max_length:
            chunk += sentence + ". "
        else:
            chunks.append(chunk.strip())
            chunk = sentence + ". "
    if chunk:
        chunks.append(chunk.strip())
    logger.debug(f"Chunked text into {len(chunks)} chunks")
    return chunks

def extract_text(file_path):
    """Extract text from supported file types."""
    extension = detect_file_type(file_path)
    logger.info(f"Extracting text from {extension} file: {file_path}")
    try:
        if extension == ".pdf":
            doc = fitz.open(file_path)
            text = "".join(page.get_text("text") for page in doc)
        elif extension == ".docx":
            doc = docx.Document(file_path)
            text = "\n".join(para.text for para in doc.paragraphs)
        elif extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif extension == ".csv":
            with open(file_path, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                text = "\n".join([", ".join(row) for row in reader])
        else:
            raise ValueError("Unsupported file for text extraction")
        logger.info(f"Extracted text of length {len(text)}")
        return text
    except Exception as e:
        logger.error(f"Error extracting text: {str(e)}")
        raise

def extract_excel_cells(file_path):
    """Extract cells from Excel file."""
    logger.info(f"Extracting cells from Excel file: {file_path}")
    try:
        wb = openpyxl.load_workbook(file_path)
        sheet = wb.active
        if not sheet:
            raise ValueError("No active sheet found")
        cells = [[cell for cell in row] for row in sheet.iter_rows(values_only=True)]
        logger.info(f"Extracted {len(cells)} rows from Excel")
        return cells
    except Exception as e:
        logger.error(f"Error extracting Excel cells: {str(e)}")
        raise

def load_model_and_tokenizer(src_lang, tgt_lang):
    """Load translation model and tokenizer, using cache if available."""
    key = (src_lang, tgt_lang)
    if key in model_cache:
        logger.info(f"Using cached model for {key}")
        return model_cache[key]
    model_name = f"Helsinki-NLP/opus-mt-{src_lang}-{tgt_lang}"
    logger.info(f"Loading model: {model_name}")
    try:
        tokenizer = MarianTokenizer.from_pretrained(model_name)
        model = MarianMTModel.from_pretrained(model_name)
        model_cache[key] = (tokenizer, model)
        logger.info(f"Model loaded: {model_name}")
        return tokenizer, model
    except Exception as e:
        logger.error(f"Error loading model {model_name}: {str(e)}")
        raise ValueError(f"Translation model for {src_lang} to {tgt_lang} not available")

def translate_text(text, src_lang, tgt_lang):
    """Translate text from source to target language."""
    logger.info(f"Translating text from {src_lang} to {tgt_lang}")
    try:
        tokenizer, model = load_model_and_tokenizer(src_lang, tgt_lang)
        chunks = chunk_text(text)
        translated_chunks = []
        for chunk in chunks:
            logger.debug(f"Translating chunk of length {len(chunk)}")
            inputs = tokenizer(chunk, return_tensors="pt", truncation=True, max_length=512)
            translated = model.generate(**inputs)
            translated_text = tokenizer.decode(translated[0], skip_special_tokens=True)
            translated_chunks.append(translated_text)
        result = " ".join(translated_chunks)
        logger.info("Text translation complete")
        return result
    except Exception as e:
        logger.error(f"Error translating text: {str(e)}")
        raise

def translate_excel_cells(cells, src_lang, tgt_lang):
    """Translate text in Excel cells."""
    logger.info(f"Translating Excel cells from {src_lang} to {tgt_lang}")
    try:
        tokenizer, model = load_model_and_tokenizer(src_lang, tgt_lang)
        translated_cells = []
        for row in cells:
            translated_row = []
            for cell in row:
                if isinstance(cell, str) and cell.strip():
                    inputs = tokenizer(cell, return_tensors="pt", truncation=True, max_length=512)
                    translated = model.generate(**inputs)
                    translated_text = tokenizer.decode(translated[0], skip_special_tokens=True)
                    translated_row.append(translated_text)
                else:
                    translated_row.append(cell)
            translated_cells.append(translated_row)
        logger.info("Excel cells translation complete")
        return translated_cells
    except Exception as e:
        logger.error(f"Error translating Excel cells: {str(e)}")
        raise

def save_translated_file(output_path, translated_text, output_format):
    """Save translated content to specified format."""
    logger.info(f"Saving translated file to {output_path} as {output_format}")
    try:
        if output_format == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(translated_text)
        elif output_format == "docx":
            doc = docx.Document()
            doc.add_paragraph(translated_text)
            doc.save(output_path)
        elif output_format == "pdf":
            c = canvas.Canvas(output_path)
            lines = translated_text.split("\n")
            y = 800
            for line in lines:
                c.drawString(40, y, line[:100])
                y -= 20
                if y < 40:
                    c.showPage()
                    y = 800
            c.save()
        elif output_format == "csv":
            reader = csv.reader(StringIO(translated_text))
            with open(output_path, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                for row in reader:
                    writer.writerow(row)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
        logger.info(f"Saved file: {output_path}")
    except Exception as e:
        logger.error(f"Error saving file: {str(e)}")
        raise

def save_translated_excel(output_path, translated_cells):
    """Save translated cells to Excel file."""
    logger.info(f"Saving translated Excel to {output_path}")
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        if not ws:
            raise ValueError("No active worksheet")
        for row_idx, row in enumerate(translated_cells, start=1):
            for col_idx, cell in enumerate(row, start=1):
                ws.cell(row=row_idx, column=col_idx, value=cell)
        wb.save(output_path)
        logger.info(f"Saved Excel: {output_path}")
    except Exception as e:
        logger.error(f"Error saving Excel: {str(e)}")
        raise

def process_file(uploaded_file, src_lang, tgt_lang, output_format):
    """Process uploaded file for translation."""
    logger.info(f"Processing file: {uploaded_file.name}, {src_lang} to {tgt_lang}, format: {output_format}")
    try:
        if not uploaded_file:
            raise ValueError("No file uploaded")

        # Save uploaded file to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp:
            tmp.write(uploaded_file.read())
            file_path = tmp.name
        logger.info(f"Created temp file: {file_path}")

        try:
            extension = detect_file_type(file_path)
            output_filename = f"{os.path.splitext(uploaded_file.name)[0]}_translated.{output_format}"
            output_path = os.path.join(OUTPUT_DIR, output_filename)

            if extension == ".xlsx":
                cells = extract_excel_cells(file_path)
                translated_cells = translate_excel_cells(cells, src_lang, tgt_lang)
                save_translated_excel(output_path, translated_cells)
            elif extension == ".csv":
                with open(file_path, "r", encoding="utf-8") as f:
                    reader = list(csv.reader(f))
                translated_cells = translate_excel_cells(reader, src_lang, tgt_lang)
                if output_format == "xlsx":
                    save_translated_excel(output_path, translated_cells)
                elif output_format == "csv":
                    with open(output_path, "w", encoding="utf-8", newline="") as f:
                        writer = csv.writer(f)
                        for row in translated_cells:
                            writer.writerow(row)
                else:
                    flat_text = "\n".join([", ".join(map(str, row)) for row in translated_cells])
                    save_translated_file(output_path, flat_text, output_format)
            else:
                text = extract_text(file_path)
                translated_text = translate_text(text, src_lang, tgt_lang)
                save_translated_file(output_path, translated_text, output_format)

            return output_path, output_filename, MIME_TYPES.get(output_format, "application/octet-stream"), "Translation successful! Download the file below."

        finally:
            if os.path.exists(file_path):
                logger.info(f"Cleaning up temp file: {file_path}")
                os.remove(file_path)

    except Exception as e:
        logger.error(f"Error processing file: {str(e)}\n{traceback.format_exc()}")
        return None, None, None, f"Translation failed: {str(e)}"

# Streamlit Interface

def main():
    st.set_page_config(page_title="LinguaLens Translation", page_icon="🌐")
    st.title("LinguaLens Translation")
    st.markdown("Upload a file, select languages, and get your translated document!")

    # File upload
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "xlsx", "csv", "txt"])

    # Language selection
    col1, col2 = st.columns(2)
    with col1:
        src_lang_name = st.selectbox("Source Language", options=list(LANGUAGES.keys()), index=0)
    with col2:
        tgt_lang_name = st.selectbox("Target Language", options=list(LANGUAGES.keys()), index=1)

    # Output format selection
    output_format = st.selectbox("Output Format", options=["txt", "docx", "pdf", "csv", "xlsx"], index=0)

    # Translate button
    if st.button("Translate"):
        if not uploaded_file:
            st.error("Please upload a file.")
        elif src_lang_name == tgt_lang_name:
            st.error("Source and target languages must be different.")
        else:
            with st.spinner("Translating..."):
                src_lang = LANGUAGES[src_lang_name]
                tgt_lang = LANGUAGES[tgt_lang_name]
                output_path, output_filename, mime_type, message = process_file(uploaded_file, src_lang, tgt_lang, output_format)
                st.write(message)
                if output_path:
                    with open(output_path, "rb") as f:
                        st.download_button(
                            label="Download Translated File",
                            data=f,
                            file_name=output_filename,
                            mime=mime_type
                        )

    # Footer
    st.markdown("---")
    st.markdown("### Powered by LinguaLens AI")
    st.markdown(
        "Built by **Adarsh Singh** | [Email](mailto:adarsh36jnp@gmail.com) | "
        "[LinkedIn](https://www.linkedin.com/in/adarsh-kumar-singh-data/)"
    )

if __name__ == "__main__":
    main()