from fastapi import FastAPI, UploadFile, Request, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from io import StringIO
from reportlab.pdfgen import canvas
from transformers import MarianMTModel, MarianTokenizer
import os
import shutil
import fitz  # PyMuPDF
import mimetypes
import openpyxl
import docx
import csv
import logging
import traceback

# Configure logging with more detailed format
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = FastAPI()

# Allow CORS with specific origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://v0-lingua-lens.vercel.app"],  # Only allow your frontend domain
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

# Create output directory if not exists
os.makedirs("translated_files", exist_ok=True)

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.error(f"Global error handler caught: {str(exc)}")
    logger.error(f"Traceback: {traceback.format_exc()}")
    return JSONResponse(
        status_code=500,
        content={"detail": f"Internal server error: {str(exc)}"}
    )

# Detect file type

def detect_file_type(file_path):
    mime_type, _ = mimetypes.guess_type(file_path)
    if mime_type and "pdf" in mime_type:
        return ".pdf"
    elif file_path.endswith(".docx"):
        return ".docx"
    elif file_path.endswith(".xlsx"):
        return ".xlsx"
    elif file_path.endswith(".csv"):
        return ".csv"
    elif file_path.endswith(".txt"):
        return ".txt"
    raise ValueError("Unsupported file type")


# Chunk long text for translation

def chunk_text(text, max_length=512):
    sentences = text.split(". ")
    chunks, chunk = [], ""
    for sentence in sentences:
        if len(chunk) + len(sentence) < max_length:
            chunk += sentence + ". "
        else:
            chunks.append(chunk.strip())
            chunk = sentence + ". "
    if chunk:
        chunks.append(chunk.strip())
    return chunks


# Text extraction

def extract_text(file_path):
    extension = detect_file_type(file_path)
    if extension == ".pdf":
        doc = fitz.open(file_path)
        return "".join(page.get_text("text") for page in doc)  # type: ignore[attr-defined]
    elif extension == ".docx":
        doc = docx.Document(file_path)
        return "\n".join(para.text for para in doc.paragraphs)
    elif extension == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif extension == ".csv":
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            return "\n".join([", ".join(row) for row in reader])
    raise ValueError("Unsupported file for text extraction")


def extract_excel_cells(file_path):
    wb = openpyxl.load_workbook(file_path)
    if not wb.sheetnames:
        raise ValueError("Excel file contains no sheets")
    sheet = wb.active
    if sheet is None:
        raise ValueError("Could not get active sheet from workbook")
    return [list(row) for row in sheet.iter_rows(values_only=True)]


# Translation

def load_model_and_tokenizer(src_lang, tgt_lang):
    model_name = f"Helsinki-NLP/opus-mt-{src_lang}-{tgt_lang}"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    return tokenizer, model


def translate_text(text, src_lang, tgt_lang):
    tokenizer, model = load_model_and_tokenizer(src_lang, tgt_lang)
    chunks = chunk_text(text)
    translated_chunks = []
    for chunk in chunks:
        inputs = tokenizer(chunk, return_tensors="pt", truncation=True, max_length=512)
        translated = model.generate(**inputs)
        translated_chunks.append(tokenizer.decode(translated[0], skip_special_tokens=True))
    return " ".join(translated_chunks)


def translate_excel_cells(cells, src_lang, tgt_lang):
    tokenizer, model = load_model_and_tokenizer(src_lang, tgt_lang)
    translated_cells = []
    for row in cells:
        translated_row = []
        for cell in row:
            if isinstance(cell, str) and cell.strip():
                inputs = tokenizer(cell, return_tensors="pt", truncation=True, max_length=512)
                translated = model.generate(**inputs)
                translated_text = tokenizer.decode(translated[0], skip_special_tokens=True)
                translated_row.append(translated_text)
            else:
                translated_row.append(cell)
        translated_cells.append(translated_row)
    return translated_cells


# Save functions

def save_translated_file(output_path, translated_text, output_format):
    if output_format == "txt":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(translated_text)
    elif output_format == "docx":
        doc = docx.Document()
        doc.add_paragraph(translated_text)
        doc.save(output_path)
    elif output_format == "pdf":
        c = canvas.Canvas(output_path)
        # Write text in chunks to fit the page
        lines = translated_text.split("\n")
        y = 800
        for line in lines:
            c.drawString(40, y, line[:100])  # Limit line length for page width
            y -= 20
            if y < 40:
                c.showPage()
                y = 800
        c.save()
    elif output_format == "csv":
        # Assume translated_text is a string with lines separated by \n and columns by ,
        reader = csv.reader(StringIO(translated_text))
        with open(output_path, "w", encoding="utf-8", newline="") as f:
            writer = csv.writer(f)
            for row in reader:
                writer.writerow(row)
    else:
        raise ValueError("Unsupported output format")


def save_translated_excel(output_path, translated_cells):
    wb = openpyxl.Workbook()
    ws = wb.active
    if ws is None:
        raise ValueError("Could not get active worksheet to save data")
    for row_idx, row in enumerate(translated_cells, start=1):
        for col_idx, cell in enumerate(row, start=1):
            ws.cell(row=row_idx, column=col_idx, value=cell)
    wb.save(output_path)


# Routes

@app.get("/", response_class=JSONResponse)
async def read_root():
    try:
        logger.info("Root endpoint called")
        return {"message": "LinguaLens Translation API is running", "status": "healthy"}
    except Exception as e:
        logger.error(f"Error in root endpoint: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload/")
async def upload_file(file: UploadFile, src_lang: str, tgt_lang: str, output_format: str = "txt"):
    try:
        logger.info(f"Upload endpoint called with file: {file.filename}, src_lang: {src_lang}, tgt_lang: {tgt_lang}")
        
        if not file.filename or not isinstance(file.filename, str):
            raise HTTPException(status_code=400, detail="Uploaded file must have a valid filename.")
        
        file_path = f"temp_{file.filename}"
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        try:
            extension = detect_file_type(file_path)
            logger.info(f"Detected file type: {extension}")
            
            if extension == ".xlsx":
                cells = extract_excel_cells(file_path)
                translated_cells = translate_excel_cells(cells, src_lang, tgt_lang)
                output_filename = f"{os.path.splitext(str(file.filename))[0]}_translated.xlsx"
                output_path = os.path.join("translated_files", output_filename)
                save_translated_excel(output_path, translated_cells)
            elif extension == ".csv":
                # Read CSV, translate each cell, and save as CSV or XLSX
                with open(file_path, "r", encoding="utf-8") as f:
                    reader = list(csv.reader(f))
                translated_cells = translate_excel_cells(reader, src_lang, tgt_lang)
                if output_format == "xlsx":
                    output_filename = f"{os.path.splitext(str(file.filename))[0]}_translated.xlsx"
                    output_path = os.path.join("translated_files", output_filename)
                    save_translated_excel(output_path, translated_cells)
                elif output_format == "csv":
                    output_filename = f"{os.path.splitext(str(file.filename))[0]}_translated.csv"
                    output_path = os.path.join("translated_files", output_filename)
                    with open(output_path, "w", encoding="utf-8", newline="") as f:
                        writer = csv.writer(f)
                        for row in translated_cells:
                            writer.writerow(row)
                else:
                    # For txt, docx, pdf: flatten to text
                    flat_text = "\n".join([", ".join(map(str, row)) for row in translated_cells])
                    output_filename = f"{os.path.splitext(str(file.filename))[0]}_translated.{output_format}"
                    output_path = os.path.join("translated_files", output_filename)
                    save_translated_file(output_path, flat_text, output_format)
            else:
                text = extract_text(file_path)
                translated_text = translate_text(text, src_lang, tgt_lang)
                output_filename = f"{os.path.splitext(str(file.filename))[0]}_translated.{output_format}"
                output_path = os.path.join("translated_files", output_filename)
                save_translated_file(output_path, translated_text, output_format)

            logger.info(f"File processed successfully. Output: {output_filename}")
            return {"message": "File translated successfully!", "download_url": f"/download/{output_filename}"}

        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

    except Exception as e:
        logger.error(f"Error in upload_file: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        if os.path.exists(file_path):
            os.remove(file_path)


@app.get("/download/{filename}")
async def download_file(filename: str):
    try:
        logger.info(f"Download endpoint called for file: {filename}")
        file_path = os.path.join("translated_files", filename)
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise HTTPException(status_code=404, detail="File not found")
        return FileResponse(file_path, filename=filename)
    except Exception as e:
        logger.error(f"Error in download_file: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health", response_class=JSONResponse)
async def health_check():
    try:
        logger.info("Health check endpoint called")
        return {"status": "healthy", "message": "API is operational"}
    except Exception as e:
        logger.error(f"Error in health check: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))
